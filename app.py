from __future__ import annotations

import os
import queue
import re
import shutil
import sys
import tempfile
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

import customtkinter as ctk
from docx import Document
from docx.text.run import Run
from tkinter import filedialog, messagebox


APP_TITLE = "Gerador de PDFs"
PLACEHOLDERS = ("__TEXTO__", "{{TEXTO}}")
RED_COLOR = "#CC092F"
OUTPUT_DIR_NAME = "resultado-pdf"
DEFAULT_DIR_NAME = "default"
MAX_TEXT_FILE_BYTES = 25 * 1024 * 1024


class AppError(Exception):
    """Erro esperado, seguro para exibir ao usuario final."""


class ProcessingCancelled(Exception):
    """Sinaliza interrupcao solicitada pelo usuario entre arquivos."""

    def __init__(self, processed_count: int) -> None:
        super().__init__("Processamento interrompido pelo usuario.")
        self.processed_count = processed_count


@dataclass(frozen=True)
class ProjectPaths:
    root_dir: Path
    default_dir: Path
    template_docx: Path
    output_dir: Path


def find_project_paths() -> ProjectPaths:
    """Localiza a raiz, a pasta default, o template unico e a saida."""
    app_dir = Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else Path(__file__).resolve().parent
    candidates = []
    for base in (Path.cwd().resolve(), app_dir):
        candidates.extend((base, base.parent))

    root_dir = None
    default_dir = None
    for candidate in candidates:
        if (candidate / DEFAULT_DIR_NAME).is_dir():
            root_dir = candidate
            default_dir = candidate / DEFAULT_DIR_NAME
            break
        if candidate.name.lower() == DEFAULT_DIR_NAME:
            root_dir = candidate.parent
            default_dir = candidate
            break

    if root_dir is None or default_dir is None:
        raise AppError(
            "A pasta 'default' nao foi encontrada. Execute o programa na raiz do projeto "
            "ou dentro da propria pasta 'default'."
        )

    templates = sorted(default_dir.glob("*.docx"))
    temporary_templates = [p for p in templates if p.name.startswith("~$")]
    templates = [p for p in templates if p not in temporary_templates]

    if not templates:
        raise AppError("Nenhum template .docx foi encontrado dentro da pasta 'default'.")
    if len(templates) > 1:
        names = ", ".join(p.name for p in templates)
        raise AppError(
            "A pasta 'default' deve conter apenas um template .docx. "
            f"Foram encontrados: {names}"
        )

    output_dir = root_dir / OUTPUT_DIR_NAME
    output_dir.mkdir(parents=True, exist_ok=True)

    return ProjectPaths(
        root_dir=root_dir,
        default_dir=default_dir,
        template_docx=templates[0],
        output_dir=output_dir,
    )


def read_text_file(path: Path) -> str:
    """Le um .txt corporativo tentando codificacoes comuns no Windows."""
    size = path.stat().st_size
    if size > MAX_TEXT_FILE_BYTES:
        limit_mb = MAX_TEXT_FILE_BYTES // (1024 * 1024)
        raise AppError(
            f"O arquivo selecionado excede o limite de {limit_mb} MB para processamento seguro."
        )

    encodings = ("utf-8-sig", "utf-8", "cp1252", "latin-1")
    last_error: Exception | None = None

    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc

    raise AppError(f"Nao foi possivel ler o arquivo '{path.name}'. Erro: {last_error}")


def clean_inserted_text(text: str) -> str:
    """Remove somente linhas 100% vazias, preservando tabs e alinhamento."""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    return "\n".join(line for line in lines if line.strip() != "")


def iter_all_paragraphs(document: Document) -> Iterable:
    """Percorre paragrafos no corpo, tabelas, cabecalhos e rodapes."""
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        containers = (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        )
        for container in containers:
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def copy_run_format(source: Run | None, target: Run) -> None:
    """Copia formatacao basica de um run sem depender de APIs privadas."""
    if source is None:
        return

    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.style = source.style

    target.font.name = source.font.name
    target.font.size = source.font.size
    target.font.bold = source.font.bold
    target.font.italic = source.font.italic
    target.font.underline = source.font.underline
    if source.font.color and source.font.color.rgb:
        target.font.color.rgb = source.font.color.rgb


def run_at_position(runs: list[Run], char_position: int) -> Run | None:
    current = 0
    fallback = runs[0] if runs else None

    for run in runs:
        text = run.text or ""
        if text:
            fallback = run
        next_position = current + len(text)
        if current <= char_position < next_position:
            return run
        current = next_position

    return fallback


def paragraph_replacement_segments(paragraph, replacement_text: str) -> list[tuple[str, Run | None]]:
    """Gera segmentos substituidos, incluindo placeholders divididos entre runs."""
    original_text = paragraph.text
    runs = list(paragraph.runs)
    segments: list[tuple[str, Run | None]] = []
    position = 0

    while position < len(original_text):
        next_match: tuple[int, str] | None = None
        for placeholder in PLACEHOLDERS:
            index = original_text.find(placeholder, position)
            if index == -1:
                continue
            if next_match is None or index < next_match[0]:
                next_match = (index, placeholder)

        if next_match is None:
            segments.append((original_text[position:], run_at_position(runs, position)))
            break

        index, placeholder = next_match
        if index > position:
            segments.append((original_text[position:index], run_at_position(runs, position)))

        segments.append((replacement_text, run_at_position(runs, index)))
        position = index + len(placeholder)

    if not segments and any(placeholder in original_text for placeholder in PLACEHOLDERS):
        segments.append((replacement_text, run_at_position(runs, 0)))

    return [(text, source_run) for text, source_run in segments if text]


def replace_placeholders_in_paragraph(paragraph, replacement_text: str) -> bool:
    """Substitui tags preservando o texto fora delas e evitando apagar runs vizinhos."""
    if not any(placeholder in paragraph.text for placeholder in PLACEHOLDERS):
        return False

    segments = paragraph_replacement_segments(paragraph, replacement_text)
    paragraph.clear()
    for text, source_run in segments:
        new_run = paragraph.add_run(text)
        copy_run_format(source_run, new_run)

    return True


def replace_placeholder(docx_path: Path, inserted_text: str, template_source: Path | None = None) -> None:
    document = Document(docx_path)
    cleaned_text = clean_inserted_text(inserted_text)
    replacements = 0

    for paragraph in iter_all_paragraphs(document):
        if replace_placeholders_in_paragraph(paragraph, cleaned_text):
            replacements += 1

    if replacements == 0:
        expected = " ou ".join(PLACEHOLDERS)
        source = template_source or docx_path
        raise AppError(
            f"Nenhuma tag de texto foi encontrada no template. Use {expected} "
            "em um paragrafo comum do Word. "
            f"Template usado: {source}"
        )

    document.save(docx_path)


def convert_docx_to_pdf_with_comtypes(docx_path: Path, pdf_path: Path) -> None:
    """Converte usando automacao COM do Word via comtypes, sem win32com."""
    try:
        import comtypes.client
    except ImportError as exc:
        raise AppError(
            "A dependencia 'comtypes' nao esta instalada. Instale com: pip install -r requirements.txt"
        ) from exc

    word = None
    document = None
    wd_alerts_none = 0
    wd_export_format_pdf = 17

    try:
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        word.DisplayAlerts = wd_alerts_none

        document = word.Documents.Open(str(docx_path), ReadOnly=False)
        try:
            document.Fields.Update()
            for section in document.Sections:
                section.Headers(1).Range.Fields.Update()
                section.Headers(2).Range.Fields.Update()
                section.Headers(3).Range.Fields.Update()
                section.Footers(1).Range.Fields.Update()
                section.Footers(2).Range.Fields.Update()
                section.Footers(3).Range.Fields.Update()
        except Exception:
            pass

        document.ExportAsFixedFormat(str(pdf_path), wd_export_format_pdf)
    except Exception as exc:
        raise AppError(
            "Nao foi possivel converter o documento para PDF. Verifique se o Microsoft Word "
            "esta instalado e se nenhum arquivo temporario esta travado em segundo plano."
        ) from exc
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def safe_pdf_name(txt_path: Path) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "_", txt_path.stem).strip()
    return f"{name or 'arquivo'}.pdf"


def check_word_available() -> tuple[bool, str]:
    """Verifica se o Microsoft Word esta disponivel via COM local."""
    try:
        import comtypes.client
    except ImportError:
        return False, "A dependencia 'comtypes' nao esta instalada."

    word = None
    try:
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        return True, "Microsoft Word disponivel."
    except Exception:
        return (
            False,
            "Microsoft Word nao foi localizado via COM. Verifique a instalacao e a licenca corporativa.",
        )
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def process_txt_files(
    txt_dir: Path,
    progress_callback: Callable[[int, int, str], None],
    cancel_event: threading.Event,
) -> int:
    paths = find_project_paths()
    txt_files = sorted(txt_dir.glob("*.txt"))

    if not txt_files:
        raise AppError("A pasta selecionada nao contem arquivos .txt.")

    processed_count = 0

    for index, txt_file in enumerate(txt_files, start=1):
        if cancel_event.is_set():
            raise ProcessingCancelled(processed_count)

        progress_callback(index, len(txt_files), f"Processando {index} de {len(txt_files)}...")

        text = read_text_file(txt_file)
        pdf_path = paths.output_dir / safe_pdf_name(txt_file)

        with tempfile.TemporaryDirectory(prefix="gerador_pdf_") as temp_dir:
            temp_docx = Path(temp_dir) / f"{txt_file.stem}.docx"
            shutil.copy2(paths.template_docx, temp_docx)
            replace_placeholder(temp_docx, text, paths.template_docx)
            if cancel_event.is_set():
                raise ProcessingCancelled(processed_count)
            convert_docx_to_pdf_with_comtypes(temp_docx, pdf_path)
            processed_count += 1

    return processed_count


class PdfGeneratorApp(ctk.CTk):
    def __init__(self) -> None:
        super().__init__()

        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        self.selected_dir: Path | None = None
        self.worker: threading.Thread | None = None
        self.events: queue.Queue[tuple[str, object]] = queue.Queue()
        self.cancel_event = threading.Event()
        self.word_available = False

        self.title(APP_TITLE)
        self.geometry("760x440")
        self.minsize(680, 380)
        self.configure(fg_color="#F4F5F7")
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

        self._build_ui()
        self.after(150, self._consume_events)
        self.after(250, self.validate_local_installation)

    def _build_ui(self) -> None:
        frame = ctk.CTkFrame(self, fg_color="#FFFFFF", corner_radius=8)
        frame.pack(fill="both", expand=True, padx=28, pady=28)

        title = ctk.CTkLabel(
            frame,
            text="Gerador de PDFs",
            font=ctk.CTkFont(size=26, weight="bold"),
            text_color="#20242A",
        )
        title.pack(anchor="w", padx=28, pady=(28, 6))

        subtitle = ctk.CTkLabel(
            frame,
            text="Converta arquivos .txt para PDF usando o template padronizado.",
            font=ctk.CTkFont(size=14),
            text_color="#5D6673",
        )
        subtitle.pack(anchor="w", padx=28, pady=(0, 24))

        self.path_label = ctk.CTkLabel(
            frame,
            text="Nenhuma pasta selecionada",
            anchor="w",
            text_color="#3B424C",
            fg_color="#F4F5F7",
            corner_radius=6,
            height=44,
        )
        self.path_label.pack(fill="x", padx=28, pady=(0, 16))

        buttons = ctk.CTkFrame(frame, fg_color="transparent")
        buttons.pack(fill="x", padx=28, pady=(0, 26))

        self.select_button = ctk.CTkButton(
            buttons,
            text="Selecionar Pasta de Arquivos (.txt)",
            fg_color=RED_COLOR,
            hover_color="#A60726",
            height=42,
            command=self.select_folder,
        )
        self.select_button.pack(side="left")

        self.generate_button = ctk.CTkButton(
            buttons,
            text="Gerar PDFs",
            fg_color="#2D333B",
            hover_color="#1E2329",
            height=42,
            width=150,
            command=self.start_generation,
            state="disabled",
        )
        self.generate_button.pack(side="left", padx=(12, 0))

        self.stop_button = ctk.CTkButton(
            buttons,
            text="Interromper",
            fg_color="#6B7280",
            hover_color="#4B5563",
            height=42,
            width=130,
            command=self.stop_generation,
            state="disabled",
        )
        self.stop_button.pack(side="left", padx=(12, 0))

        self.progress = ctk.CTkProgressBar(frame, progress_color=RED_COLOR)
        self.progress.pack(fill="x", padx=28, pady=(0, 12))
        self.progress.set(0)

        self.status_label = ctk.CTkLabel(
            frame,
            text="Pronto para iniciar.",
            anchor="w",
            text_color="#5D6673",
        )
        self.status_label.pack(fill="x", padx=28)

        footer = ctk.CTkFrame(frame, fg_color="transparent")
        footer.pack(fill="x", padx=28, pady=(28, 20), side="bottom")

        self.open_results_button = ctk.CTkButton(
            footer,
            text="Abrir pasta de resultados",
            fg_color="transparent",
            hover_color="#EEF0F3",
            text_color="#5D6673",
            border_width=1,
            border_color="#D7DBE0",
            height=34,
            width=190,
            command=self.open_results_folder,
        )
        self.open_results_button.pack(side="right")

    def select_folder(self) -> None:
        selected = filedialog.askdirectory(title="Selecione a pasta com arquivos .txt")
        if not selected:
            return
        self.selected_dir = Path(selected)
        self.path_label.configure(text=str(self.selected_dir))
        self.generate_button.configure(state="normal" if self.word_available else "disabled")
        if self.word_available:
            self.status_label.configure(text="Pasta selecionada. Clique em Gerar PDFs.")
        else:
            self.status_label.configure(text="Microsoft Word indisponivel. Geracao bloqueada.")

    def start_generation(self) -> None:
        if self.worker and self.worker.is_alive():
            return
        if self.selected_dir is None:
            messagebox.showwarning(APP_TITLE, "Selecione uma pasta com arquivos .txt antes de continuar.")
            return
        if not self.word_available:
            messagebox.showerror(
                APP_TITLE,
                "Microsoft Word nao esta disponivel via COM. A geracao de PDFs foi bloqueada.",
            )
            return

        self.progress.set(0)
        self.cancel_event.clear()
        self.status_label.configure(text="Preparando processamento...")
        self.select_button.configure(state="disabled")
        self.generate_button.configure(state="disabled")
        self.stop_button.configure(state="normal")
        self.open_results_button.configure(state="disabled")

        self.worker = threading.Thread(target=self._worker_run, daemon=False)
        self.worker.start()

    def stop_generation(self) -> None:
        if not self.worker or not self.worker.is_alive():
            return
        self.cancel_event.set()
        self.stop_button.configure(state="disabled")
        self.status_label.configure(text="Interrompendo apos a conversao atual...")

    def open_results_folder(self) -> None:
        try:
            paths = find_project_paths()
            os.startfile(paths.output_dir)
        except AppError as exc:
            messagebox.showerror(APP_TITLE, str(exc))
        except Exception as exc:
            messagebox.showerror(APP_TITLE, f"Nao foi possivel abrir a pasta de resultados: {exc}")

    def validate_local_installation(self) -> None:
        self.status_label.configure(text="Validando Microsoft Word local...")
        available, message = check_word_available()
        self.word_available = available
        if available:
            self.status_label.configure(text="Pronto para iniciar.")
            self.generate_button.configure(state="normal" if self.selected_dir else "disabled")
            return

        self.status_label.configure(text="Microsoft Word indisponivel. Geracao bloqueada.")
        self.generate_button.configure(state="disabled")
        messagebox.showerror(APP_TITLE, message)

    def on_closing(self) -> None:
        if self.worker and self.worker.is_alive():
            messagebox.showwarning(
                APP_TITLE,
                "Aguarde o termino da conversao antes de fechar. Isso evita deixar o Word aberto em segundo plano.",
            )
            return
        self.destroy()

    def _worker_run(self) -> None:
        assert self.selected_dir is not None

        def progress(current: int, total: int, status: str) -> None:
            self.events.put(("progress", (current, total, status)))

        try:
            total = process_txt_files(self.selected_dir, progress, self.cancel_event)
            self.events.put(("done", total))
        except ProcessingCancelled as exc:
            self.events.put(("cancelled", exc.processed_count))
        except AppError as exc:
            self.events.put(("error", str(exc)))
        except Exception as exc:
            self.events.put(("error", f"Ocorreu um erro inesperado: {exc}"))

    def _consume_events(self) -> None:
        try:
            while True:
                event, payload = self.events.get_nowait()
                if event == "progress":
                    current, total, status = payload
                    self.progress.set(current / total)
                    self.status_label.configure(text=status)
                elif event == "done":
                    self.progress.set(1)
                    self.status_label.configure(text=f"Concluido. {payload} PDF(s) gerado(s).")
                    messagebox.showinfo(APP_TITLE, f"{payload} PDF(s) gerado(s) com sucesso.")
                    self._unlock_buttons()
                elif event == "cancelled":
                    self.status_label.configure(text=f"Interrompido. {payload} PDF(s) gerado(s).")
                    messagebox.showinfo(APP_TITLE, f"Processamento interrompido. {payload} PDF(s) gerado(s).")
                    self._unlock_buttons()
                elif event == "error":
                    self.status_label.configure(text="Processamento interrompido.")
                    messagebox.showerror(APP_TITLE, str(payload))
                    self._unlock_buttons()
        except queue.Empty:
            pass
        finally:
            self.after(150, self._consume_events)

    def _unlock_buttons(self) -> None:
        self.select_button.configure(state="normal")
        self.generate_button.configure(state="normal" if self.selected_dir and self.word_available else "disabled")
        self.stop_button.configure(state="disabled")
        self.open_results_button.configure(state="normal")


if __name__ == "__main__":
    app = PdfGeneratorApp()
    app.mainloop()
